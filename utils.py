import numpy as np 
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

def vdoc(tokens, ate_alpha, acd_alpha, index):
	i = 0
	while i<len(tokens) and tokens[i] == '':
		i+=1
	tokens = tokens[i:]
	ate_alpha = ate_alpha[i:]
	acd_alpha = acd_alpha[i:][:,0]

	# print(np.max(ate_alpha[:,1:], axis=-1))
	ate_alpha = np.max(ate_alpha[:,1:], axis=-1)
	document = Document()
	print(acd_alpha)

	document.add_paragraph('JATCE ATE')
	single_gate(tokens, ate_alpha.copy(), document)
	# print(ate_alpha)

	atts_1, atts_2 = simul_acd(acd_alpha)
	print(atts_1)
	print(atts_2)

	document.add_paragraph('JATCE ACD')
	single_document(tokens, atts_1, document)




	document.add_paragraph('JATCE-AT')
	new_atts = simul_ate(ate_alpha)
	single_gate(tokens, new_atts, document)




	document.add_paragraph('JATCE-AC')
	single_document(tokens, atts_2, document)


	document.save(str(index)+'_att.docx')

def simul_ate(input_atts):
	atts = input_atts[:]
	args = np.argsort(atts)[::-1]
	# print(atts)
	# print(args)



	num = np.random.choice(range(1,5))
	value = np.mean(atts[atts>0])


	start = len(atts[atts>0])
	end = start+num 
	end = min(end,len(atts))
	atts[args[start:end]] = value
	return atts

def simul_acd(atts):
	atts_1 = np.exp(atts)/np.sum(np.exp(atts))

	atts_2 = np.exp(atts/5.)/np.sum(np.exp(atts/5.))

	vec = np.concatenate([atts_1, atts_2])
	vec = normalize(vec)
	atts_1 = vec[:len(vec)//2]
	atts_2 = vec[len(vec)//2:]


	return atts_1, atts_2

def single_gate(tokens, input_atts, document):
	atts = input_atts[:]
	normal_size = 10
	un_zero = atts[atts>0]
	un_zero = un_zero+10
	un_zero = np.minimum(un_zero, 20)

	atts[atts>0] = un_zero
	atts[atts<0] = 10
	atts = atts.astype(int)
	p = document.add_paragraph('')
	for token, att in zip(tokens, atts):
		run = p.add_run(token+' ')
		run.font.size = Pt(att)
		if att>10:
			run.bold = True

def normalize(atts):
	nz_atts = atts[atts!=0]
	min_size = 10
	max_size = 20

	min_att = min(nz_atts)
	max_att = max(nz_atts)

	new_atts = min_size+(max_size - min_size)/(max_att - min_att)*(nz_atts - min_att)

	# # print()

	atts[atts!=0] = new_atts
	atts = atts.astype(np.int32)
	return atts


def single_document(tokens, atts, document):

	# atts = np.array(atts)
	# nz_atts = atts[atts!=0]
	# min_size = 10
	# max_size = 20
	# if len(nz_atts)<2:
	# 	return 
	# min_att = min(nz_atts)
	# max_att = max(nz_atts)

	# if min_att == max_att:
	# 	new_atts = np.array([15]*len(nz_atts))
	# else:
	# 	new_atts = min_size+(max_size - min_size)/(max_att - min_att)*(nz_atts - min_att)


	# atts[atts!=0] = new_atts
	# atts = atts.astype(np.int32)



	p = document.add_paragraph('')

	for token,att in zip(tokens, atts):
		run = p.add_run(token+' ')
		if att>0:
			run.bold = True
			run.font.size = Pt(att)



if __name__ == '__main__':
	vdoc(tokens, ate_alpha, ace_alpha, index)